"""Update the human-facing product introduction for the current 会记 v2 build.

The checked-in product introduction used to describe the previous Whisper,
speaker diarization, template and split-settings concepts.  This script keeps
the original document's page setup and header/footer, then rebuilds the body
around the current Qwen paragraph architecture and current UI.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = ROOT / "会记_产品介绍.docx"
DEFAULT_OUTPUT = ROOT / "会记_产品介绍.docx"

FONT_NAME = "Microsoft YaHei"
INK = "172233"
MUTED = "5D6B78"
BLUE = "0B4F8A"
BRIGHT_BLUE = "2B75B5"
TEAL = "0F766E"
GOLD = "A56A00"
TABLE_HEADER = "E7EEF6"
TABLE_ALT = "F7FAFC"


def _font_name(font, name: str = FONT_NAME) -> None:
    font.name = name
    rpr = font._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), name)


def _set_run(run, *, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    _font_name(run.font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style(style, *, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    _font_name(style.font)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def _configure_styles(document: Document) -> None:
    styles = document.styles
    _set_style(styles["Normal"], size=10.6, color=INK)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.2

    _set_style(styles["Title"], size=34, color=BLUE, bold=True)
    styles["Title"].paragraph_format.space_after = Pt(8)
    _set_style(styles["Subtitle"], size=17, color=MUTED, italic=True)
    styles["Subtitle"].paragraph_format.space_after = Pt(14)
    _set_style(styles["Kicker"], size=9.5, color=TEAL, bold=True)
    styles["Kicker"].paragraph_format.space_before = Pt(4)
    styles["Kicker"].paragraph_format.space_after = Pt(8)
    _set_style(styles["Small Note"], size=8.8, color=MUTED)
    styles["Small Note"].paragraph_format.space_after = Pt(4)

    _set_style(styles["Heading 1"], size=20, color=BRIGHT_BLUE, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(7)
    styles["Heading 1"].paragraph_format.keep_with_next = True
    _set_style(styles["Heading 2"], size=14.2, color=BLUE, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)
    styles["Heading 2"].paragraph_format.keep_with_next = True
    _set_style(styles["List Bullet"], size=10.6, color=INK)
    styles["List Bullet"].paragraph_format.space_after = Pt(3)
    styles["List Bullet"].paragraph_format.line_spacing = 1.15
    _set_style(styles["List Number"], size=10.6, color=INK)
    styles["List Number"].paragraph_format.space_after = Pt(3)
    styles["List Number"].paragraph_format.line_spacing = 1.15


def _clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _paragraph(document: Document, text: str = "", *, style: str = "Normal", color: str | None = None, size: float | None = None, bold: bool = False, italic: bool = False, align=None, after: float | None = None):
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    if text:
        run = paragraph.add_run(text)
        _set_run(run, color=color, size=size, bold=bold if bold else None, italic=italic if italic else None)
    return paragraph


def _heading(document: Document, text: str, level: int = 1):
    paragraph = document.add_heading(text, level=level)
    if level == 1 and (text.startswith("1.") or text.startswith("2.")):
        paragraph.paragraph_format.page_break_before = True
    for run in paragraph.runs:
        _set_run(run, color=BRIGHT_BLUE if level == 1 else BLUE, size=20 if level == 1 else 14.2, bold=True)
    return paragraph


def _rich_paragraph(document: Document, parts: list[tuple[str, dict]], *, style: str = "Normal", after: float | None = None):
    paragraph = document.add_paragraph(style=style)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    for text, options in parts:
        run = paragraph.add_run(text)
        _set_run(run, **options)
    return paragraph


def _bullet(document: Document, text: str) -> None:
    _paragraph(document, text, style="List Bullet")


def _number(document: Document, text: str) -> None:
    _paragraph(document, text, style="List Number")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, *, color: str = "D4DEE9", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_margins(cell, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _write_cell(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 9.7) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    _set_run(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _table(document: Document, headers: list[str], rows: list[list[str]], *, widths: list[float] | None = None, font_size: float = 9.7):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for idx, cell in enumerate(table.rows[0].cells):
        _set_cell_shading(cell, TABLE_HEADER)
        _set_cell_border(cell)
        _set_cell_margins(cell)
        _write_cell(cell, headers[idx], bold=True, color=BLUE, size=font_size)
        if widths:
            cell.width = Inches(widths[idx])
    _prevent_row_split(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        _prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            _set_cell_border(cell)
            _set_cell_margins(cell)
            if len(table.rows) % 2 == 0:
                _set_cell_shading(cell, TABLE_ALT)
            _write_cell(cell, row_values[idx], bold=idx == 0, color=INK, size=font_size)
            if widths:
                cell.width = Inches(widths[idx])
    return table


def _callout(document: Document, label: str, text: str, *, fill: str = "E8F3F1", border: str = "0F766E", label_color: str = TEAL):
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_cell_border(cell, color=border, size="10")
    _set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.18
    label_run = paragraph.add_run(label + "  ")
    _set_run(label_run, size=10.1, color=label_color, bold=True)
    body_run = paragraph.add_run(text)
    _set_run(body_run, size=10.1, color=INK)
    _prevent_row_split(table.rows[0])
    return table


def _page_break(document: Document) -> None:
    # Numbered Heading 1 paragraphs carry their own page break.  Keeping this
    # helper as a no-op avoids a blank page when the preceding content already
    # ends at the bottom of a page.
    return None


def build_document(source: Path, output: Path) -> Path:
    document = Document(source)
    _clear_body(document)
    _configure_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    document.core_properties.title = "会记 v2 产品介绍（当前版本）"
    document.core_properties.subject = "Qwen 双模型、段落式实时转写、翻译与会议行动闭环"
    document.core_properties.keywords = "会记,实时会议,Qwen3-ASR,段落转写,会议纪要,To-do-list"

    # Cover and current-version summary.
    _paragraph(document, "产品介绍  ·  当前版本说明", style="Kicker", align=WD_ALIGN_PARAGRAPH.LEFT)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("会记")
    _set_run(run, size=34, color=BLUE, bold=True)
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("本地优先的实时会议记录工作台")
    for run in subtitle.runs:
        _set_run(run, size=17, color=MUTED, italic=True)
    _rich_paragraph(document, [("让会议从“有人记录”变为“可实时理解、可复盘、可执行”。", {"size": 12, "color": INK, "bold": True})], after=12)

    _table(document, ["项目", "当前产品定义"], [
        ["产品形态", "Windows 本地运行的实时会议记录工作台"],
        ["当前架构", "Qwen3-ASR 双模型 + 连续语音段落输出，两套实时识别架构可选"],
        ["语言能力", "中文、英文、德文；中文方言由小千问判断并按官方类别归一化记录"],
        ["交付结果", "实时段落、原文与简体中文翻译、会议纪要、To-do-list 和本次会议文件"],
        ["数据边界", "音频、ASR、VAD 和翻译优先在本机执行；纪要与 To-do 依赖配置的 Jimo SSE 节点"],
    ], widths=[1.35, 5.85], font_size=9.8)
    _callout(document, "一句话定位", "会记把本地录音变成连续、可回看、可执行的会议记录：录音时先看实时段落，停录后完成翻译，再生成纪要和行动项。")
    _heading(document, "产品一览", level=1)
    _table(document, ["实时理解", "会后可用", "本地优先", "结果可追溯"], [[
        "按连续语音段落持续更新原文和译文",
        "翻译完成后生成会议纪要与 To-do-list",
        "Qwen、VAD、OPUS-MT 在本机运行",
        "schema 2.0、revision 和可下载文件",
    ]], widths=[1.8, 1.8, 1.8, 1.8], font_size=9.1)

    _page_break(document)

    # Positioning and users.
    _heading(document, "1. 产品定位与目标用户", level=1)
    _paragraph(document, "会记面向需要边开会边理解、会后快速复盘并形成行动清单的个人和小团队。它不是单纯的录音机，也不是只在会后给出一段摘要的黑盒，而是一条从音频采集、实时段落、翻译到纪要和 To-do 的本地化处理链。")
    _heading(document, "典型场景", level=2)
    _bullet(document, "跨语种项目会议：中文、英文、德文混合讨论，实时查看原文，英文和德文同步补充简体中文翻译。")
    _bullet(document, "研发评审与项目例会：保留连续语音段落和时间信息，减少会后回听和手工整理。")
    _bullet(document, "访谈、销售和客户沟通：先按段落复盘原始表达，再从已保存内容生成纪要和可追踪行动项。")
    _bullet(document, "对隐私敏感的本地会议：音频、识别和翻译留在本机；只有纪要与 To-do 所需文本才交给已配置的 Jimo 节点。")

    _heading(document, "产品承诺", level=2)
    _table(document, ["用户价值", "产品表现"], [
        ["实时可见", "WebSocket 持续更新当前段落；partial 只修订当前卡片，不让同一句话重复堆叠。"],
        ["模型透明", "界面选择的是两套实时识别架构，清楚说明主模型、语言/方言判断模型和回退关系。"],
        ["本地优先", "Qwen ASR、VAD 与英文/德文翻译采用本地模型；外部服务只承担配置范围内的纪要和 To-do。"],
        ["可复盘", "同一段落使用 segment_id 和 revision 追踪更新，并保存 JSON、JSONL、Markdown 等结果文件。"],
        ["可恢复", "翻译失败保留原文并支持重试；停录后按固定顺序完成保存、翻译和纪要前置状态。"],
    ], widths=[1.35, 5.85], font_size=9.6)
    _callout(document, "当前版本边界", "会记不保存人员身份，不做说话人身份分离，也不存在会后再跑一套旧 ASR 的隐藏阶段；产品以连续段落和原文/译文为核心记录单位。", fill="FFF7E8", border=GOLD, label_color=GOLD)

    _page_break(document)

    # Current model architecture.
    _heading(document, "2. 当前核心能力", level=1)
    _heading(document, "2.1 两套实时识别架构", level=2)
    _paragraph(document, "当前界面的“实时识别架构”选择的是实时 ASR 的主模型和回退关系，不是“全程只使用一种模型”。两套架构都会使用 Qwen3-ASR-0.6B 做语言/方言判断；录音开始后，本次会议的架构会锁定。")
    _table(document, ["界面选项", "实际链路", "适合场景", "录音开始后"], [
        ["大+小千问 · 质量优先", "1.7B 主识别；0.6B 负责语言/方言判断并作为实时回退", "更看重实时识别质量和稳定性", "架构锁定，异常时切换到 0.6B"],
        ["小千问主识别 · 低延迟", "0.6B 主识别；0.6B 仍负责语言/方言判断，1.7B 作为异常回退", "更看重响应速度和较低资源占用", "架构锁定，异常时切换到 1.7B"],
    ], widths=[1.55, 2.55, 1.7, 1.4], font_size=8.9)
    _callout(document, "模型选择说明", "“大+小千问”表示 1.7B 主识别 + 0.6B 判断/兜底；“小千问主识别”表示 0.6B 主识别 + 1.7B 兜底，并不代表全程只使用大千问或小千问。", fill="EAF2FC", border=BRIGHT_BLUE, label_color=BLUE)

    _heading(document, "2.2 语言、方言与翻译", level=2)
    _bullet(document, "当前结果保留 zh / en / de / unknown；中文方言由 0.6B 判断，并按 Qwen 官方 22 类归一化记录，包含浙江、粤语香港/广东口音、吴语和闽南语等代表类别。")
    _bullet(document, "普通话和中文方言直接规范为简体中文；英文、德文以及中英混合段落按稳定原文前缀排队翻译为简体中文。")
    _bullet(document, "语言首次出现先作为候选，连续确认后才触发段落边界；短暂 unknown 不会制造额外段落。")

    _heading(document, "2.3 实时处理链路", level=2)
    _table(document, ["阶段", "当前行为"], [
        ["音频输入", "浏览器采集 16 kHz、单声道 PCM16，通过 WebSocket 发送到本机服务。"],
        ["分段与识别", "VAD 和静音阈值决定连续语音段；0.6B 判断语言/方言，主 ASR 按已选架构输出 partial。"],
        ["段落聚合", "同一连续语音段使用同一个 segment_id；partial 持续修订，段落关闭后形成稳定原文。"],
        ["翻译与呈现", "英文/德文按稳定前缀进入串行翻译队列，前端通过 paragraph_update 原地更新段落卡片。"],
    ], widths=[1.25, 6.0], font_size=9.4)

    _page_break(document)

    # User experience and settings.
    _heading(document, "3. 一次会议的完整体验", level=1)
    _heading(document, "3.1 从新建到导出", level=2)
    _number(document, "新建会议并填写名称；服务检查完成后进入当前会议工作台。")
    _number(document, "在录音前打开“录音与识别设置”，选择实时识别架构，确认输入设备和背景声过滤。")
    _number(document, "开始录音后，实时转写区按连续段落更新原文；英文和德文在对应段落下补充简体中文译文。")
    _number(document, "停录后，系统依次完成最后音频 flush、实时 ASR、翻译队列和本地保存，并将会议推进到可生成纪要的状态。")
    _number(document, "生成会议纪要后，系统从已保存内容提取 To-do-list；失败时可分别重试，不影响已保存的录音和转写。")
    _number(document, "在右侧结果区查看会议纪要、To-do-list 和本次文件，并下载 Markdown、JSON 等结果。")

    _heading(document, "3.2 当前界面与设置设计", level=2)
    _table(document, ["界面区域", "当前能力"], [
        ["左侧栏", "会记品牌、会议历史、搜索和新建会议；深浅色主题切换为品牌区域右上角的纯图标按钮。"],
        ["中间工作区", "录音状态、计时、设备与环境音、实时转写段落；可回到最新段落。"],
        ["右侧结果区", "会议纪要、To-do-list 和本次文件，随会议状态显示生成、重试和下载操作。"],
        ["录音与识别设置", "模型架构、输入设备、背景声过滤、静音结束阈值、实时刷新间隔、录音分段长度和录音保留。"],
    ], widths=[1.65, 5.6], font_size=9.4)
    _callout(document, "设置简化", "当前设置已经合并为一个页面，不再区分“普通设置”和“高级设置”，也不再保留“参数模板”概念。点击设置弹窗外部会直接取消并关闭，不保存本次未提交的修改。", fill="E8F3F1", border=TEAL, label_color=TEAL)
    _callout(document, "主题切换", "浅色/深色按钮不显示文字，只使用月亮/太阳图标；按钮位于左侧品牌区域右上角，主题偏好保存在本机浏览器。", fill="EAF2FC", border=BRIGHT_BLUE, label_color=BLUE)

    _page_break(document)

    # Reliability and storage.
    _heading(document, "4. 稳定性、结果与恢复", level=1)
    _heading(document, "4.1 段落式结果契约", level=2)
    _paragraph(document, "当前版本以“连续语音段落”而不是逐句消息作为前端和存储单位。一个段落在实时过程中可以被多次修订，但始终由同一个 segment_id 标识；revision 用于追踪顺序，前端收到 paragraph_update 后原地更新，不追加重复节点。")
    _table(document, ["机制", "作用"], [
        ["paragraph_update", "统一承载实时段落的原文、译文、语言/方言、模型和 closed 状态。"],
        ["稳定前缀翻译", "只翻译已经稳定的原文前缀，过期 source_revision 结果不会覆盖更新后的内容。"],
        ["串行队列与重试", "翻译任务按段落顺序执行；持续失败保留原文，并可调用 translation/retry。"],
        ["schema 2.0", "transcript.json 保存当前 paragraphs，JSONL 追加 revision 事件；旧会议不迁移。"],
    ], widths=[1.75, 5.5], font_size=9.5)

    _heading(document, "4.2 停录顺序与会议状态", level=2)
    _callout(document, "固定顺序", "flush 最后音频段 → 等待实时 ASR → 等待翻译队列 → recording_state=complete → ready_for_summary。只有完成翻译队列后，右侧才开放生成纪要和 To-do-list。", fill="FFF7E8", border=GOLD, label_color=GOLD)
    _bullet(document, "实时 ASR 运行期间保持模型链路，停录不再触发一套隐藏的会后 Whisper/说话人处理流程。")
    _bullet(document, "翻译失败不抹掉原文；会议结果仍可查看，修复服务或模型后可以重新排队翻译。")
    _bullet(document, "模型缺失、服务异常和处理失败通过明确状态展示，避免把失败伪装成成功。")

    _heading(document, "4.3 本次会议输出", level=2)
    _table(document, ["文件", "用途"], [
        ["transcript.json / transcript.jsonl", "当前段落投影与追加修订事件，schema_version=2.0。"],
        ["meeting_transcript.md", "按段落整理的会议原文，包含时间、语言/方言和段落内容。"],
        ["translated_zh.md", "英文、德文及混合段落对应的简体中文译文。"],
        ["original_zh.md / original_en.md / original_de.md", "按语言拆分的原文文件，便于复盘和二次处理。"],
        ["manifest.json / audio_manifest.json / session_state.json", "会议元数据、音频清单和生命周期状态。"],
        ["audio/（可选）", "按照录音保留设置保存的本地音频分段。"],
    ], widths=[2.4, 4.85], font_size=9.3)

    _page_break(document)

    # Privacy and deployment.
    _heading(document, "5. 数据边界与部署方式", level=1)
    _heading(document, "5.1 本地优先，但边界透明", level=2)
    _table(document, ["处理内容", "当前归属"], [
        ["音频采集、VAD、Qwen ASR、语言/方言判断", "本机浏览器与本机服务；默认使用本地模型缓存。"],
        ["英文/德文 → 简体中文翻译", "本地 OPUS-MT 模型，不依赖云端翻译 API。"],
        ["会议纪要与 To-do-list", "使用本地配置的 Jimo SSE 节点；prompt 只使用段落时间、语言/方言、原文和译文。"],
        ["身份信息与说话人分离", "当前版本不保存人员身份，也不提供说话人编号或说话人重排。"],
    ], widths=[2.6, 4.65], font_size=9.3)
    _callout(document, "隐私边界", "会记是本地优先的会议记录工作台，不把“本地 ASR/翻译”和“纪要服务调用”混写成同一条链路；服务级地址、授权和模型下载策略由部署配置控制。", fill="E8F3F1", border=TEAL, label_color=TEAL)

    _heading(document, "5.2 运行要求", level=2)
    _table(document, ["项目", "当前要求"], [
        ["操作系统", "Windows 10/11 x64；建议使用 Chrome 或 Edge。"],
        ["运行时", "Python 3.11；FastAPI 本机服务；浏览器通过 WebSocket 传输 16 kHz 单声道 PCM16。"],
        ["模型准备", "Qwen3-ASR-1.7B、Qwen3-ASR-0.6B、FunASR FSMN-VAD 和 OPUS-MT en→zh / de→zh。"],
        ["生产建议", "设置 MEETING_ASR_AUTODOWNLOAD=0 和 MEETING_TRANSLATION_AUTODOWNLOAD=0，只读取已准备好的本地模型缓存。"],
    ], widths=[1.45, 5.8], font_size=9.3)

    _heading(document, "5.3 当前服务配置", level=2)
    _table(document, ["配置", "当前值 / 说明"], [
        ["MEETING_ASR_PRIMARY", "Qwen/Qwen3-ASR-1.7B"],
        ["MEETING_ASR_FALLBACK", "Qwen/Qwen3-ASR-0.6B"],
        ["MEETING_ASR_LANGUAGE_ID", "Qwen/Qwen3-ASR-0.6B"],
        ["TRANSCRIPT_SCHEMA_VERSION", "2.0"],
        ["language_id_min_seconds", "0.8 秒左右开始语言/方言判断"],
    ], widths=[2.75, 4.5], font_size=9.5)

    _page_break(document)

    # Version notes and source of truth.
    _heading(document, "6. 当前版本变更摘要", level=1)
    _paragraph(document, "本版产品介绍已按会记 v2 的当前实现重新整理，重点修正了模型、结果契约和界面描述之间的不一致。")
    _table(document, ["已更新项", "当前版本描述"], [
        ["模型链路", "移除旧 Whisper、large-v3、Resemblyzer 和会后第二套 ASR 的产品描述，统一改为 Qwen3-ASR 双模型架构。"],
        ["模型选择文案", "明确区分“质量优先”和“低延迟”两套架构，说明主识别、语言/方言判断与 fallback，不再暗示全程只用一种模型。"],
        ["设置页面", "删除参数模板概念，合并普通/高级设置为一个页面，保留核心稳定性参数和本次会议生效范围。"],
        ["交互行为", "设置弹窗点击外部取消且不保存；深浅色切换使用无文字图标，并放在左侧品牌区域右上角。"],
        ["存储与展示", "以连续段落和 paragraph_update 为核心，使用 transcript schema 2.0，支持原文、译文、纪要、To-do 和文件下载。"],
    ], widths=[1.45, 5.8], font_size=9.2)
    _callout(document, "阅读提示", "产品能力以当前代码和配置为准：README.md、DEPLOYMENT.md、realtime_meeting/config.py、runtime.py、session.py、storage.py 以及 web 目录是本版本的实现依据。", fill="EAF2FC", border=BRIGHT_BLUE, label_color=BLUE)

    _heading(document, "验收重点", level=2)
    _bullet(document, "中英德切换、中文方言代表类别、短暂 unknown、连续长语音和中英混合段落。")
    _bullet(document, "同一 segment_id 的多次 partial 修订、稳定前缀翻译、过期结果丢弃和失败重试。")
    _bullet(document, "停录顺序、schema 2.0 文件输出、会议纪要和 To-do 生成前置状态。")
    _bullet(document, "浅色/深色主题、设置单页、弹窗外部取消不保存、模型选择文案和输入文字可读性。")
    _paragraph(document, "文档版本：会记 v2 · 当前版本说明。", style="Small Note", after=0)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description="Update the current 会记 product introduction DOCX")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    print(build_document(args.input, args.output))


if __name__ == "__main__":
    main()
