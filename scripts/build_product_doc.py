"""Build a short product brief for the single-model Qwen architecture.

The product brief is deliberately generated from the current contract instead
of carrying historical implementation details.  ``python-docx`` remains an
optional documentation-only dependency; the runtime does not need it.
"""

from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - documentation helper
    raise SystemExit("build_product_doc.py requires python-docx") from exc


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "meeting_product_brief_v2.docx"


def _style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Microsoft YaHei"
    styles["Title"].font.size = Pt(28)
    styles["Heading 1"].font.name = "Microsoft YaHei"
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 2"].font.name = "Microsoft YaHei"
    styles["Heading 2"].font.size = Pt(12.5)


def _paragraph(document: Document, text: str, *, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold


def _bullets(document: Document, values: list[str]) -> None:
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def _table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value


def build_document() -> Path:
    document = Document()
    _style(document)
    document.core_properties.title = "实时会议转写产品说明（单模型 Qwen 段落架构）"
    document.core_properties.subject = "Qwen 单模型、同模语言确认、段落式实时转写与翻译"

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("实时会议转写")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Qwen 1.7B 单模型 + 同模语言确认 + 段落式输出").bold = True

    document.add_heading("1. 产品定义", level=1)
    _paragraph(document, "系统在录音期间持续输出可读的会议段落：唯一常驻的 Qwen3-ASR-1.7B 同时负责最终 ASR、分段级语言确认和语言冲突重识别；英文和德文通过本地 OPUS-MT 翻译为简体中文，中文及中文方言直接规范为简体中文。")
    _bullets(document, [
        "一个连续的语言/方言语音段对应一个段落卡片，partial 只修订当前卡片。",
        "静音达到阈值或语言稳定切换时结束段落；时间戳音频帧用于按确认时刻重切片，技术性音频切分不会制造额外显示段落。",
        "系统不保存人员身份、手动语言锁定或会后第二套 ASR 结果；本地会后复译默认关闭。",
        "旧会议数据不迁移；新存储契约使用 transcript schema 2.0。",
    ])

    document.add_heading("2. 模型与适配层", level=1)
    _table(document, ["用途", "模型", "传参策略"], [
        ["实时 ASR、语言确认与冲突重识别", "Qwen/Qwen3-ASR-1.7B", "同一 checkpoint 负责最终识别、约 1 秒后的分段语言探测和冲突重识别；en=English、de=German，中文及方言使用对应上下文提示。"],
        ["兼容回退配置", "仍归一到 Qwen/Qwen3-ASR-1.7B", "旧 fallback/LID 配置字段保留以兼容 API，但单模型模式不加载 0.6B，也不切换第二个 ASR。"],
        ["英文/德文翻译", "OPUS-MT en→zh / de→zh", "稳定前缀异步翻译；按 segment_id 合并任务，source_revision 过期结果丢弃；服务启动时 warm-up。"],
        ["VAD", "FunASR FSMN-VAD", "保留既有音频阈值与静音结束机制，并保留带时间戳的音频帧。"],
    ])

    document.add_heading("3. 实时数据流", level=1)
    _paragraph(document, "浏览器 PCM 音频 → VAD/时间帧保留/技术分段 → Qwen3-ASR-1.7B partial/final → 同模语言证据聚合与时间边界重切片 → 段落聚合器 → segment_id 合并翻译队列 → TranscriptStore 与 paragraph_update → 前端段落卡片。")
    _bullets(document, [
        "约 1 秒后由同一 Qwen 模型探测语言；新语言需要连续 3 次一致证据才切换，短暂 unknown 被忽略，中文方言只更新 speech_variant。",
        "语言切换确认后按时间戳切开旧/新窗口再整理 ASR，避免一个 chunk 混入两种语言造成漏字、重复或错序。",
        "partial 采用约 1 秒的 latest-wins；final 优先；翻译任务按 segment_id 合并，旧 source_revision 返回后不能覆盖新文本。",
        "翻译使用本地 OPUS-MT 并在启动时 warm-up；会后本地复译默认关闭，后续如有质量需求可单独请求外部翻译智能体。",
        "停止流程：flush 最后音频段 → 等待实时 ASR → 等待实时翻译队列 → 跳过默认关闭的会后复译 → recording_state=complete → ready_for_summary。",
    ])

    document.add_heading("4. 存储和 WebSocket 契约", level=1)
    _paragraph(document, "每个段落由同一个 segment_id 标识，通过 revision 追加修订事件。段落记录包含 id、segment_id、start、end、language、speech_variant、language_confidence、text、translation_zh、translation_status、revision、source_revision、closed、asr_model、language_source、translation_model。")
    _paragraph(document, "实时事件统一为 paragraph_update；前端按 segment_id 原地更新，不追加逐句节点。transcript.json 保存 paragraphs，transcript.jsonl 保存追加事件，schema_version 为 2.0。")
    _table(document, ["接口", "用途"], [
        ["GET /api/v2/meetings/{id}", "会议状态、活动段落、语言和翻译队列快照。"],
        ["GET /api/v2/meetings/{id}/transcript", "schema 2.0 段落列表。"],
        ["POST /api/v2/meetings/{id}/translation/retry", "重新排队失败的段落翻译。"],
        ["POST /api/v2/meetings/{id}/summary", "用户手动触发一次三段结果生成。"],
    ])

    document.add_heading("5. 导出和验收重点", level=1)
    _bullets(document, [
        "meeting_transcript.md、translated_zh.md 和语言原文文件按段落输出时间、语言/方言、原文和译文。",
        "manifest.json 不含人员或旧后台处理字段；不再生成 speaker_segments.json。",
        "验收覆盖中英德切换、普通话与 Qwen 官方 22 类中文方言代表类别切换、长连续语音、短暂低音量、中英夹杂、噪声和多人同时讲话。",
        "测试验证同一 segment_id 的多次 partial 更新、稳定前缀的过期结果丢弃、中文不调用翻译模型、技术切分合并和停止顺序。",
    ])

    document.add_heading("6. 当前模型配置", level=1)
    _table(document, ["配置", "值"], [
        ["MEETING_SINGLE_ASR_MODEL", "1"],
        ["MEETING_ASR_PRIMARY", "Qwen/Qwen3-ASR-1.7B"],
        ["MEETING_ASR_FALLBACK", "Qwen/Qwen3-ASR-1.7B（兼容别名，不加载第二模型）"],
        ["MEETING_ASR_LANGUAGE_ID", "Qwen/Qwen3-ASR-1.7B（兼容别名，同模确认）"],
        ["MEETING_PARTIAL_INTERVAL_MS", "1000"],
        ["MEETING_LANGUAGE_ID_MIN_SECONDS", "1.0"],
        ["MEETING_LANGUAGE_CONFLICT_\nCONFIRMATIONS", "3"],
        ["MEETING_POST_TRANSLATION_ENABLED", "0（默认关闭本地会后复译）"],
        ["TRANSCRIPT_SCHEMA_VERSION", "2.0"],
    ])
    _paragraph(document, "模型能力与部署前检查以仓库 README.md、DEPLOYMENT.md、config.py、runtime.py、session.py 和 storage.py 为准。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
